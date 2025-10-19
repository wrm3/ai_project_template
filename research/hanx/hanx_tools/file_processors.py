import pandas as pd
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import pdfplumber
from typing import List, Dict, Union, Optional
import csv
import json
import os

class WordProcessor:
    @staticmethod
    def read_docx(file_path: str) -> str:
        """Read content from a Word document."""
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    @staticmethod
    def create_docx(file_path: str, content: Union[str, List[str]], 
                   title: Optional[str] = None) -> None:
        """Create a Word document with the given content."""
        doc = Document()
        if title:
            doc.add_heading(title, 0)

        if isinstance(content, str):
            doc.add_paragraph(content)
        else:
            for paragraph in content:
                doc.add_paragraph(paragraph)

        doc.save(file_path)

    @staticmethod
    def add_to_docx(file_path: str, content: Union[str, List[str]]) -> None:
        """Add content to an existing Word document."""
        doc = Document(file_path)
        if isinstance(content, str):
            doc.add_paragraph(content)
        else:
            for paragraph in content:
                doc.add_paragraph(paragraph)
        doc.save(file_path)

class ExcelProcessor:
    @staticmethod
    def read_excel(file_path: str, sheet_name: Optional[str] = None) -> pd.DataFrame:
        """Read data from an Excel file."""
        return pd.read_excel(file_path, sheet_name=sheet_name)

    @staticmethod
    def create_excel(file_path: str, data: Union[pd.DataFrame, Dict], 
                    sheet_name: str = 'Sheet1') -> None:
        """Create an Excel file with the given data."""
        if isinstance(data, dict):
            df = pd.DataFrame(data)
        else:
            df = data
        df.to_excel(file_path, sheet_name=sheet_name, index=False)

    @staticmethod
    def append_excel(file_path: str, data: Union[pd.DataFrame, Dict], 
                    sheet_name: str = 'Sheet1') -> None:
        """Append data to an existing Excel file."""
        if isinstance(data, dict):
            df = pd.DataFrame(data)
        else:
            df = data
        
        with pd.ExcelWriter(file_path, mode='a', engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name=sheet_name, index=False)

class CSVProcessor:
    @staticmethod
    def read_csv(file_path: str, delimiter: str = ',') -> List[Dict]:
        """Read data from a CSV file."""
        with open(file_path, 'r', newline='', encoding='utf-8') as file:
            reader = csv.DictReader(file, delimiter=delimiter)
            return list(reader)

    @staticmethod
    def create_csv(file_path: str, data: List[Dict], 
                  fieldnames: Optional[List[str]] = None) -> None:
        """Create a CSV file with the given data."""
        if not fieldnames and data:
            fieldnames = list(data[0].keys())

        with open(file_path, 'w', newline='', encoding='utf-8') as file:
            writer = csv.DictWriter(file, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(data)

    @staticmethod
    def append_csv(file_path: str, data: List[Dict]) -> None:
        """Append data to an existing CSV file."""
        with open(file_path, 'a', newline='', encoding='utf-8') as file:
            writer = csv.DictWriter(file, fieldnames=list(data[0].keys()))
            writer.writerows(data)

class HTMLProcessor:
    @staticmethod
    def read_html(file_path: str) -> str:
        """Read content from an HTML file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    @staticmethod
    def parse_html(content: str) -> BeautifulSoup:
        """Parse HTML content."""
        return BeautifulSoup(content, 'html.parser')

    @staticmethod
    def create_html(file_path: str, content: str, title: str = '') -> None:
        """Create an HTML file with the given content."""
        html_template = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{title}</title>
        </head>
        <body>
            {content}
        </body>
        </html>
        """
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(html_template)

class PDFProcessor:
    @staticmethod
    def read_pdf(file_path: str) -> List[str]:
        """Read content from a PDF file."""
        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text())
        return text

    @staticmethod
    def extract_tables(file_path: str) -> List[List[List]]:
        """Extract tables from a PDF file."""
        tables = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables.extend(page.extract_tables())
        return tables

    @staticmethod
    def extract_images(file_path: str, output_dir: str) -> List[str]:
        """Extract images from a PDF file."""
        image_files = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                for j, image in enumerate(page.images):
                    image_path = os.path.join(output_dir, f'image_p{i+1}_{j+1}.png')
                    with open(image_path, 'wb') as img_file:
                        img_file.write(image['stream'].get_data())
                    image_files.append(image_path)
        return image_files 